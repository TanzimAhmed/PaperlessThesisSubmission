from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.conf import settings
import os


class CoverPage:
    def __init__(self, course_title, section, title, instructor, group_name, group_members):
        self.course_title = course_title
        self.section = section
        self.title = title
        self.instructor = instructor
        self.group_name = group_name
        self.group_members = group_members
        self.document = None
        self.table = None

    def generate_page(self):
        self.document = Document()

        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
        font.color.rgb = RGBColor(89, 89, 89)

        self.add_paragraph('NORTH SOUTH UNIVERSITY')
        self.add_logo()

        paragraph = self.add_paragraph(self.course_title.upper())
        run = paragraph.runs[0]
        # run.add_tab() run.add_break()
        run = paragraph.add_run('  Section ' + self.section)
        run.font.size = Pt(22)

        self.add_paragraph(self.title.title(), bold=False)
        self.add_paragraph('', 12, False)
        paragraph = self.add_paragraph('Instructor', font_size=18)
        run = paragraph.runs[0]
        run.add_break()
        run = paragraph.add_run(self.instructor.upper())
        run.font.size = Pt(18)
        run.font.name = 'Calibri Light'

        self.add_paragraph('', 12, False)
        paragraph = self.add_paragraph('Due Date', font_size=18)
        run = paragraph.runs[0]
        run.add_break()
        run = paragraph.add_run('27-th November 2019, Wednesday')
        run.font.size = Pt(18)
        run.font.name = 'Calibri Light'

        self.add_group_info(self.group_name)

        section = self.document.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)

        destination = os.path.join(settings.BASE_DIR, 'static/files/test_cover_page.docx')
        print(self.document.paragraphs)
        self.document.save(destination)

    def add_paragraph(self, text, font_size=36, bold=True):
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        return paragraph

    def add_group_info(self, group_name, table_rows=3, table_cols=2):
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run('Group: ')
        run.font.size = Pt(18)
        run.font.name = 'Calibri Light'
        run = paragraph.add_run(group_name)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = 'Calibri Light'
        self.table = self.document.add_table(rows=table_rows, cols=table_cols)
        row = 0
        col = 0
        for member in self.group_members:
            if row == table_rows:
                row = 0
                col = 1
                self.add_group_members(row, col, member['name'].upper(), member['id'])
            else:
                self.add_group_members(row, col, member['name'].upper(), member['id'])
            row += 1

    def add_group_members(self, row, col, student_name, student_id, font_size=16):
        cell = self.table.rows[row].cells[col]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(student_name)
        run.font.size = Pt(font_size)
        run.font.name = 'Calibri Light'
        run.add_break()
        run = paragraph.add_run(student_id)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = 'Calibri Light'
        return paragraph

    def add_logo(self, width=1.5):
        paragraph = self.document.add_paragraph()
        source = os.path.join(settings.BASE_DIR, 'static/images/NSU_logo.png')
        paragraph.add_run().add_picture(source, width=Inches(width))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


class DocumentTest:
    def __init__(self, file_name):
        self.document = Document(file_name)
        self.sections = self.document.sections

    def check_properties(self):
        i = 0
        print("Printing Page Properties")
        for section in self.sections:
            i += 1
            print(f'\n\nFor section: {i}')
            print(f'Start Type: {section.start_type}')
            print(f'Orientation: {section.orientation}')
            print(f'Page Height: {section.page_height.inches} inch(es)')
            print(f'Page Width: {section.page_width.inches} inch(es)')
            print(f'Top Margin: {section.top_margin.inches} inch(es)')
            print(f'Right Margin: {section.right_margin.inches} inch(es)')
            print(f'Left Margin: {section.left_margin.inches} inch(es)')
            print(f'Bottom Margin: {section.bottom_margin.inches} inch(es)')

    def check_paragraph_styles(self):
        paragraphs = self.document.paragraphs
        print('\nPrinting Paragraph Properties')
        print(f'Paragraphs Present: {len(paragraphs)}')
        for paragraph in paragraphs:
            paragraph_format = paragraph.paragraph_format
            print(f'Paragraph Alignment: {paragraph_format.alignment}')
            print(f'Paragraph Line-Spacing: {paragraph_format.line_spacing} inch(es)')
            self.check_fonts(paragraph)

    def check_fonts(self, paragraph):
        runs = paragraph.runs
        print('\nPrinting Run Properties')
        print(f'Runs Present: {len(runs)}')
        for run in runs:
            print(f'Font Name: {run.font.name}')
            print(f'Font Size: {run.font.size}')
            print(f'Font Is Bold: {run.font.bold}')
            print(f'Font Color: {run.font.color}')


if __name__ == '__main__':
    members = [
        {
            'name': 'sumiya noorjahan',
            'id': '151 2065 642'
        }, {
            'name': 'fareha alamgir',
            'id': '152 1464 642'
        }, {
            'name': 'fahad bin bari shovo',
            'id': '161 0162 042'
        }, {
            'name': 'shadat hossain pabel',
            'id': '161 2332 042'
        }, {
            'name': 'tanzim al din ahmed',
            'id': '162 1203 042'
        }
    ]

    cover_page = CoverPage('cse 499a', '21', 'paperless thesis submission', 'dr. md shahriar karim', members)
    cover_page.generate_page()
    print('Page Generated')
