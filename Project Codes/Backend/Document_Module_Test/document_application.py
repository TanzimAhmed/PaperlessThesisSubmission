from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

document = Document()

document.add_heading('Test Document', level=0)

text = 'Lorem ipsum dolor sit amet, consectetur adipiscing elit. Phasellus tristique enim vel ligula tristique, ' \
       'ac dictum mi vulputate. Aenean vel justo a libero convallis auctor in sed leo. Integer rutrum lorem quis ' \
       'nisl lacinia hendrerit. Etiam sed fringilla turpis. Nullam ut cursus dolor, id sodales enim. Vestibulum ' \
       'faucibus aliquet tristique. Vivamus eu nisl magna. Aliquam nec est eu nisl tempor commodo sit amet dapibus ' \
       'dolor. Cras vestibulum pellentesque erat, sagittis lobortis lacus efficitur non. Donec fermentum urna quis ' \
       'nibh eleifend, sit amet dignissim tellus lacinia. Sed fermentum nibh vitae ligula vestibulum, maximus ornare' \
       ' urna luctus.'

style = document.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)

paragraph = document.add_paragraph(text)
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


# Adding and Centering a picture
paragraph = document.add_paragraph()
paragraph.add_run().add_picture('NSU_logo.png', width=Inches(2))
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

"""
# Alternate way of center aligning picture
picture = document.paragraphs[-1]
picture.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
"""

paragraph = document.add_paragraph('Some more paragraph ' + text)
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
document.save('text_document.docx')

print('Document saved!')
