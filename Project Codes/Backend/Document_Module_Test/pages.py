from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import fitz
import os


class CoverPage:
    def __init__(self, course_title, section, title, instructor, group_name, group_members):
        self.course_title = course_title
        self.section = section
        self.title = title
        self.instructor = instructor
        self.group_name = group_name
        self.group_members = group_members
        self.document = None
        self.table = None

    def generate_page(self):
        self.document = Document()

        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
        font.color.rgb = RGBColor(89, 89, 89)

        self.add_paragraph('NORTH SOUTH UNIVERSITY')
        self.add_logo()

        paragraph = self.add_paragraph(self.course_title.upper())
        run = paragraph.runs[0]
        # run.add_tab() run.add_break()
        run = paragraph.add_run('  Section ' + self.section)
        run.font.size = Pt(22)

        self.add_paragraph(self.title.title(), bold=False)
        self.add_paragraph('', 12, False)
        paragraph = self.add_paragraph('Instructor', font_size=18)
        run = paragraph.runs[0]
        run.add_break()
        run = paragraph.add_run(self.instructor.upper())
        run.font.size = Pt(18)
        run.font.name = 'Calibri Light'

        self.add_paragraph('', 12, False)
        paragraph = self.add_paragraph('Due Date', font_size=18)
        run = paragraph.runs[0]
        run.add_break()
        run = paragraph.add_run('27-th November 2019, Wednesday')
        run.font.size = Pt(18)
        run.font.name = 'Calibri Light'

        self.add_group_info(self.group_name)

        section = self.document.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)

        destination = 'test_cover_page.docx'
        print(self.document.paragraphs)
        self.document.save(destination)

    def add_paragraph(self, text, font_size=36, bold=True):
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        return paragraph

    def add_group_info(self, group_name, table_rows=3, table_cols=2):
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run('Group: ')
        run.font.size = Pt(18)
        run.font.name = 'Calibri Light'
        run = paragraph.add_run(group_name)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = 'Calibri Light'
        self.table = self.document.add_table(rows=table_rows, cols=table_cols)
        row = 0
        col = 0
        for member in self.group_members:
            if row == table_rows:
                row = 0
                col = 1
                self.add_group_members(row, col, member['name'].upper(), member['id'])
            else:
                self.add_group_members(row, col, member['name'].upper(), member['id'])
            row += 1

    def add_group_members(self, row, col, student_name, student_id, font_size=16):
        cell = self.table.rows[row].cells[col]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(student_name)
        run.font.size = Pt(font_size)
        run.font.name = 'Calibri Light'
        run.add_break()
        run = paragraph.add_run(student_id)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = 'Calibri Light'
        return paragraph

    def add_logo(self, width=1.5):
        paragraph = self.document.add_paragraph()
        source = 'NSU_logo.png'
        paragraph.add_run().add_picture(source, width=Inches(width))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


class DocumentTest:
    def __init__(self, file_name):
        self.document = Document(file_name)
        self.sections = self.document.sections
        self.paragraphs = self.document.paragraphs
        self.page_properties = None
        self.text_properties = None
        self.base_properties = {
            'page_height': {11.69},
            'page_width': {8.27},
            'top_margin': {1.0},
            'right_margin': {1.0},
            'left_margin': {1.0},
            'bottom_margin': {1.0},
            'line_spacing': {1.15, 1.0},
            'font_name': {'Cambria'}
        }
        self.errors = None
        self.base_properties['line_spacing'].add(None)
        self.base_properties['font_name'].add(None)

    def get_page_properties(self):
        properties = {
            'length': len(self.sections),
            'start_type': set(),
            'orientation': set(),
            'page_height': set(),
            'page_width': set(),
            'top_margin': set(),
            'right_margin': set(),
            'left_margin': set(),
            'bottom_margin': set()
        }
        for section in self.sections:
            properties['start_type'].add(section.start_type)
            properties['orientation'].add(section.orientation)
            properties['page_height'].add(round(section.page_height.inches, 2))
            properties['page_width'].add(round(section.page_width.inches, 2))
            properties['top_margin'].add(round(section.top_margin.inches, 2))
            properties['right_margin'].add(round(section.right_margin.inches, 2))
            properties['left_margin'].add(round(section.left_margin.inches, 2))
            properties['bottom_margin'].add(round(section.bottom_margin.inches, 2))
        self.page_properties = properties
        return self.page_properties

    def get_text_properties(self):
        properties = {
            'length': len(self.document.paragraphs),
            'alignment': set(),
            'line_spacing': set(),
            'style_name': set(),
            'font_name': set(),
            'font_size': set()
        }
        for paragraph in self.paragraphs:
            style_name = paragraph.style.name
            properties['alignment'].add(self.document.styles[style_name].paragraph_format.alignment)
            properties['line_spacing'].add(self.document.styles[style_name].paragraph_format.line_spacing)
            properties['style_name'].add(style_name)
            properties['font_name'].update(self.get_fonts(paragraph))
            if self.document.styles[style_name].font.size is not None:
                properties['font_size'].add(self.document.styles[style_name].font.size.pt)
        self.text_properties = properties
        return self.text_properties

    def get_fonts(self, paragraph):
        fonts = set()
        runs = paragraph.runs
        style_name = paragraph.style.name
        for run in runs:
            if run.font.name is None:
                fonts.add(self.document.styles[style_name].font.name)
            fonts.add(run.font.name)
        return list(fonts)

    def is_valid_format(self):
        self.get_page_properties()
        self.get_text_properties()
        errors = []
        valid = True
        if not self.page_properties['page_height'] == self.base_properties['page_height']:
            errors.append(self.error_message('page height'))
            valid = False
        if not self.page_properties['page_width'] == self.base_properties['page_width']:
            errors.append(self.error_message('page width'))
            valid = False
        if not self.page_properties['top_margin'] == self.base_properties['top_margin']:
            errors.append(self.error_message('top margin'))
            valid = False
        if not self.page_properties['right_margin'] == self.base_properties['right_margin']:
            errors.append(self.error_message('right margin'))
            valid = False
        if not self.page_properties['left_margin'] == self.base_properties['left_margin']:
            errors.append(self.error_message('left margin'))
            valid = False
        if not self.page_properties['bottom_margin'] == self.base_properties['bottom_margin']:
            errors.append(self.error_message('bottom margin'))
            valid = False
        if not self.text_properties['line_spacing'] <= self.base_properties['line_spacing']:
            errors.append(self.error_message('line spacing present'))
            valid = False
        if not self.text_properties['font_name'] <= self.base_properties['font_name']:
            errors.append(self.error_message('fonts present'))
            valid = False
        if not valid:
            self.errors = errors
        return valid

    def error_message(self, item):
        return f'Invalid {item}'

    def check_heading_styles(self):
        pass

    def check_properties(self):
        styles = self.document.styles
        i = 0
        print("Printing Page Properties")
        """
        for style in styles:
            print(f"Document Style name: {style.name}")
            print(f"Document Style element: {style.element}")
            print(f"Font: {styles['Heading 1'].font.name}")
        """
        for section in self.sections:
            i += 1
            print(f'\n\nFor section: {i}')
            print(f'Start Type: {section.start_type}')
            print(f'Orientation: {section.orientation}')
            print(f'Page Height: {section.page_height.inches} inch(es)')
            print(f'Page Width: {section.page_width.inches} inch(es)')
            print(f'Top Margin: {section.top_margin.inches} inch(es)')
            print(f'Right Margin: {section.right_margin.inches} inch(es)')
            print(f'Left Margin: {section.left_margin.inches} inch(es)')
            print(f'Bottom Margin: {section.bottom_margin.inches} inch(es)')

    def check_paragraph_styles(self):
        paragraphs = self.document.paragraphs
        print('\nPrinting Heading Paragraph Properties')
        print(f'Paragraphs Present: {len(paragraphs)}')
        for paragraph in paragraphs:
            style_name = paragraph.style.name
            print(f'Paragraph Alignment: {self.document.styles[style_name].paragraph_format.alignment}')
            print(f'Paragraph Line-Spacing: {self.document.styles[style_name].paragraph_format.line_spacing} inch(es)')
            print(f'Paragraph style: {style_name}')
            print(f'Font Name: {self.document.styles[style_name].font.name}')
            if self.document.styles[style_name].font.size is not None:
                print(f'Font Size: {self.document.styles[style_name].font.size.pt}')
            print()
            # self.check_fonts(paragraph)

    def check_fonts(self, paragraph):
        runs = paragraph.runs
        print('\nPrinting Run Properties')
        print(f'Runs Present: {len(runs)}')
        for run in runs:
            print(f'Font Name: {run.font.name}')
            print(f'Font Size: {run.font.size}')
            print(f'Font Is Bold: {run.font.bold}')
            print(f'Font Color: {run.font.color}')


class PdfDocumentTest:
    def __init__(self, name):
        self.document = fitz.Document(name)
        self.page_count = self.document.pageCount
        self.fonts = None
        self.margin = None
        self.line_height = None
        self.page_height = None
        self.page_width = None
        self.errors = None
        self.base_properties = {
            'page_height': 841.68,
            'page_width': 595.44,
            'top_margin': 72,
            'right_margin': 72,
            'left_margin': 72,
            'bottom_margin': 72,
            'line_spacing': [1.15, 1.0],
            'font_name': ['Calibri', 'Cambria']
        }

    def print_properties(self):
        print(self.document.metadata)
        print(self.document.pageCount)
        print(self.document.getToC())
        fonts = set()
        x0, x1, y0, y1 = [], [], [], []
        for page_no in range(self.page_count):
            page = self.document.loadPage(page_no)
            print(f'Bounds; CropBox; MediaBox:{page_no + 1}; {page.bound()}; {page.CropBox}; {page.MediaBox}; {page.rect}')

            font_list = page.getFontList()
            for font in font_list:
                fonts.add(font[3])
            # print(f'Links: {page.getLinks()}')
            blocks = page.getText('blocks')
            for block in blocks:
                if 791.5 < round(block[3], 2) < 807:
                    # print(block)
                    continue
                if round(block[3], 2) > 750:
                    print(block)
                x0.append(round(block[0], 2))
                x1.append(round(block[2], 2))
                y0.append(round(block[1], 2))
                y1.append(round(block[3], 2))
        print('X0, X1, Y0, Y1 :')
        print(min(x0))
        print(max(x1))
        print(min(y0))
        print(max(y1))
        print(f'Fonts: {fonts}')

    def is_valid_format(self):
        dimensions, justified, fonts, margins = None, None, None, None
        for page_no in range(self.page_count):
            page = self.document.loadPage(page_no)
            dimensions = self.check_page_dimensions(page_no, page.bound())
            data = page.getText('dict')
            block_no = 1
            for block in data['blocks']:
                if block['type'] != 1:
                    justified = self.is_justified(page_no + 1, block_no, block['lines'])
                    fonts = self.check_fonts(page_no, block_no, block['lines'])
                    self.get_margins(block['bbox'])
                    block_no += 1
        margins = self.check_margins()
        if dimensions and justified and fonts and margins:
            return True
        else:
            return False

    def is_justified(self, page_no, block_no, lines):
        if self.line_height is None:
            self.line_height = 0
        valid = True
        for i in range(len(lines) - 2):
            if page_no in [1, 5, 6, 7]:
                continue
            text = ''
            for span in lines[i]['spans']:
                text += span['text']
            if len(text) < 70:
                continue
            line_height_min, line_height_max = 5000, 0
            line_height_min = min(line_height_min, lines[i+1]['bbox'][1] - lines[i]['bbox'][3])
            line_height_max = max(line_height_max, lines[i+1]['bbox'][1] - lines[i]['bbox'][3])
            line_height = (line_height_max + line_height_min) / 2
            self.line_height = round((self.line_height + line_height) / 2, 2)
            if not (lines[i]['bbox'][0] < 72.5 and lines[i]['bbox'][2] > 522.5):
                valid = False
        if not valid:
            if self.errors is None:
                self.errors = []
            self.errors.append(f'Line Not Justified: Page: {page_no}, Paragraph: {block_no}')
        return valid

    def check_fonts(self, page_no, block_no, lines):
        if self.fonts is None:
            self.fonts = set()
        font = None
        valid = True
        for line in lines:
            for span in line['spans']:
                if span['text'].isspace():
                    continue
                font = span['font'].split('+')[-1]
                font = font.split('-')[0]
                self.fonts.add(font)
                if font not in self.base_properties['font_name']:
                    valid = False
        if not valid:
            if self.errors is None:
                self.errors = []
            self.errors.append(f'Contains illegal Font: {font}; Page: {page_no}, Paragraph: {block_no}')
        return valid

    def get_margins(self, coordinate):
        if self.margin is None:
            self.margin = {
                'left': 5000,
                'right': 0,
                'top': 5000,
                'bottom': 0
            }
        valid = True
        if 791.5 < round(coordinate[3], 2) < 807:
            return
        self.margin['left'] = round(min(self.margin['left'], coordinate[0]), 2)
        self.margin['right'] = round(max(self.margin['right'], coordinate[2]), 2)
        self.margin['top'] = round(min(self.margin['top'], coordinate[1]), 2)
        self.margin['bottom'] = round(max(self.margin['bottom'], coordinate[3]), 2)

    def check_margins(self):
        valid = True
        errors = []
        self.margin['right'] = round(self.base_properties['page_width'] - self.margin['right'], 2)
        self.margin['bottom'] = round(self.base_properties['page_height'] - self.margin['bottom'], 2)
        if not self.base_properties['left_margin'] - 5 < self.margin['left'] < self.base_properties['left_margin'] + 5:
            valid = False
            errors.append(f'Contains illegal Left Margin.')
        if not self.base_properties['right_margin'] - 5 < self.margin['right'] \
                < self.base_properties['right_margin'] + 5:
            valid = False
            errors.append(f'Contains illegal Right Margin.')
        if not self.base_properties['top_margin'] - 5 < self.margin['top'] < self.base_properties['top_margin'] + 5:
            valid = False
            errors.append(f'Contains illegal Top Margin.')
        if not self.base_properties['bottom_margin'] - 5 < self.margin['bottom'] \
                < self.base_properties['bottom_margin'] + 5:
            valid = False
            errors.append(f'Contains illegal Bottom Margin.')
        if not valid:
            if self.errors is None:
                self.errors = []
            for error in errors:
                self.errors.append(error)
        return valid

    def check_page_dimensions(self, page_no, coordinates):
        errors = []
        valid = True
        self.page_height = round(coordinates[3], 2)
        self.page_width = round(coordinates[2], 2)
        if not self.base_properties['page_height'] - 5 < coordinates[3] < self.base_properties['page_height'] + 5:
            valid = False
            errors.append(f'Contains illegal Page Height; Page: {page_no}')
        if not self.base_properties['page_width'] - 5 < coordinates[2] < self.base_properties['page_width'] + 5:
            valid = False
            errors.append(f'Contains illegal Page Width; Page: {page_no}')
        if not valid:
            if self.errors is None:
                self.errors = []
            for error in errors:
                self.errors.append(error)
        return valid


if __name__ == '__main__':
    members = [
        {
            'name': 'sumiya noorjahan',
            'id': '151 2065 642'
        }, {
            'name': 'fareha alamgir',
            'id': '152 1464 642'
        }, {
            'name': 'fahad bin bari shovo',
            'id': '161 0162 042'
        }, {
            'name': 'shadat hossain pabel',
            'id': '161 2332 042'
        }, {
            'name': 'tanzim al din ahmed',
            'id': '162 1203 042'
        }
    ]

    cover_page = CoverPage('cse 499a', '21', 'paperless thesis submission', 'dr. md shahriar karim', members)
    cover_page.generate_page()
    print('Page Generated')
